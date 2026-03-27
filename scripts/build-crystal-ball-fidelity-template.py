from __future__ import annotations

"""
Deprecated.

This script used fixed python-docx paragraph indices to inject placeholders into a
template derived from the source document. That approach is no longer authoritative.

The project now uses:
1. a semantic template profile (`retail-crystal-ball-v1`) for structure and density
2. a clean DOCX shell plus runtime XML fill / cleanup / package repair

Keep this file only as a historical reference of the legacy path. Do not use it as
the production template builder.
"""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
VENDOR = ROOT / ".reqagent" / "vendor" / "python-docxtpl"
if str(VENDOR) not in sys.path:
    sys.path.insert(0, str(VENDOR))

from docx import Document  # type: ignore


SOURCE = ROOT / ".reqagent" / "workspaces" / "ws_77420f3f-2cde-47b7-8b0c-c03abe621356-636865ed732e" / "docs" / "(零售水晶球三期)202502新增代发管理-非标准化代发业务单位管理需求说明书V1.1_20250106.docx"
OUTPUT = ROOT / ".reqagent" / "workspaces" / "ws_77420f3f-2cde-47b7-8b0c-c03abe621356-636865ed732e" / "docs" / "零售业务需求说明书_精细模板.docx"


def copy_run_style(source_run, target_run):
    if source_run is None:
        return
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.style = source_run.style
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    if source_run.font.color is not None and source_run.font.color.rgb is not None:
        target_run.font.color.rgb = source_run.font.color.rgb
    try:
        if source_run.font.highlight_color is not None:
            target_run.font.highlight_color = source_run.font.highlight_color
    except ValueError:
        pass


def replace_paragraph(paragraph, text: str):
    first_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    copy_run_style(first_run, run)


def set_cell_text(cell, text: str):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "")


def replace_range(doc: Document, start: int, end: int, prefix: str):
    counter = 1
    for index in range(start, end + 1):
        text = doc.paragraphs[index].text.strip()
        if not text:
            continue
        replace_paragraph(doc.paragraphs[index], f"{{{{{prefix}{counter}}}}}")
        counter += 1


def main():
    raise SystemExit(
        "Deprecated legacy script. Use the DOCX runtime profile/fill pipeline instead."
    )

    doc = Document(str(SOURCE))

    # Cover and metadata
    replace_paragraph(doc.paragraphs[2], "{{项目名称}}需求说明书")
    replace_paragraph(doc.paragraphs[9], "制作单位：{{制作单位}}")
    replace_paragraph(doc.paragraphs[10], "文档版本号：{{文档版本号}}")
    replace_paragraph(doc.paragraphs[11], "日期：{{日期}}")
    replace_paragraph(doc.paragraphs[13], "编写人员：{{编写人员}}")
    replace_paragraph(doc.paragraphs[21], "校对人员：{{校对人员}}")
    replace_paragraph(doc.paragraphs[31], "{{签署日期}}")

    # Template-only business content
    replace_paragraph(doc.paragraphs[77], "{{需求背景}}")
    replace_paragraph(doc.paragraphs[79], "{{业务目标}}")
    replace_paragraph(doc.paragraphs[81], "{{业务价值}}")
    replace_paragraph(doc.paragraphs[82], "")
    replace_paragraph(doc.paragraphs[84], "{{术语内容}}")
    replace_paragraph(doc.paragraphs[85], "")
    replace_paragraph(doc.paragraphs[88], "{{业务概述}}")
    replace_paragraph(doc.paragraphs[89], "")
    replace_paragraph(doc.paragraphs[90], "")
    replace_paragraph(doc.paragraphs[92], "{{业务处理流程说明}}")
    replace_paragraph(doc.paragraphs[93], "")
    replace_paragraph(doc.paragraphs[94], "")
    replace_paragraph(doc.paragraphs[95], "")
    replace_paragraph(doc.paragraphs[96], "")
    replace_paragraph(doc.paragraphs[98], "{{现状问题概述}}")
    replace_paragraph(doc.paragraphs[99], "{{同业现状}}")
    replace_paragraph(doc.paragraphs[101], "{{现存问题}}")
    replace_paragraph(doc.paragraphs[120], "业务功能一：{{功能名称1}}")

    replace_range(doc, 128, 140, "业务流程说明")
    replace_paragraph(doc.paragraphs[143], "SR_{{功能编号1}}")
    replace_range(doc, 144, 145, "功能详述")
    replace_range(doc, 147, 157, "业务规则")

    replace_paragraph(doc.paragraphs[198], "是否涉及使用外部数据：{{外部数据是否使用}}；外部数据审核编号：{{审核编号}}")
    replace_paragraph(doc.paragraphs[201], "是否含有客户相关信息：{{外部数据含客信息}}")
    replace_paragraph(doc.paragraphs[203], "授权方式描述：{{授权方式}}")
    replace_paragraph(doc.paragraphs[206], "是否影响监管报送：{{监管报送影响}}")
    replace_paragraph(doc.paragraphs[207], "如不影响，请详细说明原因：{{不影响原因}}")
    replace_paragraph(doc.paragraphs[210], "如影响，请写明涉及报送系统：{{系统名称}}")
    replace_paragraph(doc.paragraphs[211], "影响表和字段范围：{{字段范围}}；逻辑调整方案：{{调整方案}}")
    replace_paragraph(doc.paragraphs[214], "是否已落实数据分级分类：{{数据分级落实}}")
    replace_paragraph(doc.paragraphs[217], "{{数据分析需求}}")
    replace_paragraph(doc.paragraphs[221], "{{非功能性需求}}")
    replace_paragraph(doc.paragraphs[224], "{{系统需求}}")
    replace_paragraph(doc.paragraphs[166], "{{支付系统需求}}")
    replace_paragraph(doc.paragraphs[169], "{{回单系统需求}}")
    replace_paragraph(doc.paragraphs[172], "{{报表需求}}")
    replace_paragraph(doc.paragraphs[175], "{{询证函需求}}")
    replace_paragraph(doc.paragraphs[178], "{{京智柜面需求}}")
    replace_paragraph(doc.paragraphs[182], "{{核算引擎需求}}")
    replace_paragraph(doc.paragraphs[185], "{{对手信息需求}}")
    replace_paragraph(doc.paragraphs[188], "{{通知类业务需求}}")
    replace_paragraph(doc.paragraphs[191], "{{联网核查需求}}")

    # Change log table
    change_log = doc.tables[0]
    headers = ["变更号", "日期", "图号/表号/段落号", "A/M/D", "题目或简短描述", "更改申请号"]
    for row_index in range(1, 3):
        for col_index, _header in enumerate(headers):
            token = f"{{{{变更{row_index}_{col_index + 1}}}}}"
            set_cell_text(change_log.rows[row_index].cells[col_index], token)

    # Department table
    dept_table = doc.tables[1]
    for row_index in range(1, len(dept_table.rows)):
        set_cell_text(dept_table.rows[row_index].cells[0], f"{{{{部门{row_index}}}}}")
        set_cell_text(dept_table.rows[row_index].cells[1], f"{{{{职责{row_index}}}}}")

    # Input elements table
    input_table = doc.tables[2]
    for row_index in range(1, len(input_table.rows)):
        n = row_index
        row = input_table.rows[row_index]
        set_cell_text(row.cells[1], f"{{{{输入字段{n}}}}}")
        set_cell_text(row.cells[2], f"{{{{输入类型{n}}}}}")
        set_cell_text(row.cells[3], f"{{{{输入必填{n}}}}}")
        set_cell_text(row.cells[4], f"{{{{输入枚举{n}}}}}")
        set_cell_text(row.cells[5], f"{{{{输入备注{n}}}}}")

    # Output elements table
    output_table = doc.tables[3]
    for row_index in range(1, len(output_table.rows)):
        n = row_index
        row = output_table.rows[row_index]
        set_cell_text(row.cells[1], f"{{{{输出字段{n}}}}}")
        set_cell_text(row.cells[2], f"{{{{输出类型{n}}}}}")
        set_cell_text(row.cells[3], f"{{{{输出必填{n}}}}}")
        set_cell_text(row.cells[4], f"{{{{输出枚举{n}}}}}")
        set_cell_text(row.cells[5], f"{{{{输出备注{n}}}}}")

    doc.save(str(OUTPUT))
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
