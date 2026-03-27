from __future__ import annotations

import shutil
import tempfile
import zipfile
from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_CANDIDATES = [
    Path(
        "/Users/dylanthomas/Library/Containers/com.tencent.WeWorkMac/Data/Documents/Profiles/"
        "2C0671F1696E8AE4EC1D64B4DBAF2B71/Caches/Files/2026-03/"
        "69131cb999fddc070f3bcf2e0a2a2736/用户需求说明书参考模板.docx"
    ),
    ROOT / "docs" / "用户需求说明书参考模板.docx",
]
OUTPUT = ROOT / "docs" / "用户需求说明书_Base_clean.docx"

HEADING_STYLE_NAMES = {
    "heading 1",
    "heading 2",
    "heading 3",
    "heading 4",
    "heading 5",
    "toc 1",
    "toc 2",
    "toc 3",
    "toc 4",
    "toc 5",
    "toc 6",
    "toc 7",
    "toc 8",
    "toc 9",
}


def resolve_source() -> Path:
    for candidate in SOURCE_CANDIDATES:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("未找到用户需求说明书参考模板")


def clear_paragraph(paragraph, text: str = "") -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag.endswith("}pPr"):
            continue
        paragraph_element.remove(child)
    if text:
        paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_table_after(document: Document, paragraph, rows: int, cols: int, style_name: str):
    table = document.add_table(rows=rows, cols=cols)
    table.style = style_name
    paragraph._p.addnext(table._tbl)
    return table


def set_cell_text(cell, value: str) -> None:
    cell.text = value


def build_function_catalog_table(document: Document) -> None:
    anchor = next((p for p in document.paragraphs if p.text.strip() == "功能分类"), None)
    if anchor is None:
        return

    table = insert_table_after(document, anchor, rows=13, cols=4, style_name="Normal Table")
    headers = ["序号", "功能模块", "功能名称", "备注"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header)

    for row_index in range(1, len(table.rows)):
        cells = table.rows[row_index].cells
        set_cell_text(cells[0], f"{{{{功能清单序号{row_index}}}}}")
        set_cell_text(cells[1], f"{{{{功能模块{row_index}}}}}")
        set_cell_text(cells[2], f"{{{{功能清单名称{row_index}}}}}")
        set_cell_text(cells[3], f"{{{{功能清单备注{row_index}}}}}")


def normalize_change_log_table(document: Document) -> None:
    if len(document.tables) < 1:
        return
    table = document.tables[0]
    placeholders = [
        ["更改号", "日期", "图号/表号/段落号", "A/M/D", "题目或简短描述", "更改申请号"],
        ["{{变更1_1}}", "{{变更1_2}}", "{{变更1_3}}", "{{变更1_4}}", "{{变更1_5}}", "{{变更1_6}}"],
    ]

    for row_index, row_values in enumerate(placeholders):
        row = table.rows[row_index]
        for cell_index, value in enumerate(row_values):
            set_cell_text(row.cells[cell_index], value)

    for row in table.rows[2:]:
        for cell in row.cells:
            set_cell_text(cell, "")


def normalize_department_table(document: Document) -> None:
    if len(document.tables) < 2:
        return
    table = document.tables[1]
    for row_index, row in enumerate(table.rows[1:], start=1):
        set_cell_text(row.cells[0], f"{{{{部门{row_index}}}}}")
        set_cell_text(row.cells[1], f"{{{{职责{row_index}}}}}")


def normalize_io_table(table, prefix: str) -> None:
    headers = ["序号", "字段名称", "类型", "是否必填", "枚举值", "备注"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header)

    for row_index, row in enumerate(table.rows[1:], start=1):
        set_cell_text(row.cells[0], f"{{{{{prefix}序号{row_index}}}}}")
        set_cell_text(row.cells[1], f"{{{{{prefix}字段{row_index}}}}}")
        set_cell_text(row.cells[2], f"{{{{{prefix}类型{row_index}}}}}")
        set_cell_text(row.cells[3], f"{{{{{prefix}必填{row_index}}}}}")
        set_cell_text(row.cells[4], f"{{{{{prefix}枚举{row_index}}}}}")
        set_cell_text(row.cells[5], f"{{{{{prefix}备注{row_index}}}}}")


def normalize_feature_tables(document: Document) -> None:
    if len(document.tables) < 4:
        return
    normalize_io_table(document.tables[2], "输入")
    normalize_io_table(document.tables[3], "输出")


def normalize_cover(document: Document) -> None:
    replacements = {
        "(项目名称)": "{{项目名称}}",
        "制作单位：(部门)": "制作单位：{{制作单位}}",
        "文档版本号:": "文档版本号：{{文档版本号}}",
        "日期：": "日期：{{日期}}",
        "编写人员:": "编写人员：{{编写人员}}",
        "校对人员:": "校对人员：{{校对人员}}",
        "年  月    日": "{{签署日期}}",
        "业务功能一：XXX": "业务功能一：{{功能名称1}}",
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            clear_paragraph(paragraph, replacements[text])


def remove_instructional_body_paragraphs(document: Document) -> None:
    first_heading_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "概述"),
        None,
    )
    if first_heading_index is None:
        return

    paragraphs_to_remove = []
    for paragraph in document.paragraphs[first_heading_index:]:
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        text = paragraph.text.strip()
        if not text:
            paragraphs_to_remove.append(paragraph)
            continue
        if style_name in HEADING_STYLE_NAMES:
            continue
        paragraphs_to_remove.append(paragraph)

    for paragraph in paragraphs_to_remove:
        remove_paragraph(paragraph)


def set_core_properties(document: Document) -> None:
    core = document.core_properties
    core.title = "用户需求说明书 Base Template"
    core.subject = "用户需求说明书 Base Template"
    core.author = "ReqAgent"
    core.comments = "Clean DOCX base shell for generic requirement document export."


def clean_package_relationships(output_path: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="reqagent-base-shell-") as temp_dir_raw:
        temp_dir = Path(temp_dir_raw)
        with zipfile.ZipFile(output_path) as archive:
            archive.extractall(temp_dir)

        word_dir = temp_dir / "word"
        rels_path = word_dir / "_rels" / "document.xml.rels"
        if rels_path.exists():
            rels_xml = rels_path.read_text("utf-8")
            kept = []
            for match in re.finditer(r'<Relationship\b[^>]*Target="([^"]+)"[^>]*/>', rels_xml):
                target = match.group(1)
                if target.startswith("media/") or target.startswith("embeddings/"):
                    continue
                kept.append(match.group(0))

            root_match = re.search(r"<Relationships\b[^>]*>", rels_xml)
            if root_match:
                rels_path.write_text(
                    f"{rels_xml[:root_match.end()]}{''.join(kept)}</Relationships>",
                    "utf-8",
                )

        for folder_name in ("media", "embeddings"):
            folder = word_dir / folder_name
            if folder.exists():
                shutil.rmtree(folder, ignore_errors=True)

        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for file_path in temp_dir.rglob("*"):
                if file_path.is_dir():
                    continue
                archive.write(file_path, file_path.relative_to(temp_dir))


def main() -> None:
    source = resolve_source()
    document = Document(source)

    normalize_cover(document)
    normalize_change_log_table(document)
    normalize_department_table(document)
    normalize_feature_tables(document)
    remove_instructional_body_paragraphs(document)
    build_function_catalog_table(document)
    set_core_properties(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    clean_package_relationships(OUTPUT)

    print(f"source={source}")
    print(f"output={OUTPUT}")


if __name__ == "__main__":
    main()
